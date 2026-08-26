"""Unit tests: DOCX insertion (run-level, no duplicate), LaTeX insertion,
bibliography update, audit gates (8A/8B)."""
import json
import sys
import unittest
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parent.parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS))

FIXTURES = Path(__file__).resolve().parent.parent.parent / "tests" / "fixtures"

from docx import Document  # noqa: E402
from insert_docx import apply_insertion_plan, locate_sentence_run  # noqa: E402
from insert_latex import parse_latex, apply_insertion_plan as latex_apply  # noqa: E402
from update_references import (  # noqa: E402
    update_bibtex_file, apa_reference_order,
)
from format_bibtex import extract_dois_from_bib  # noqa: E402
from audit_entailment import audit_links, _heuristic_verdict  # noqa: E402
from audit_bibliography import audit_document_integrity  # noqa: E402


class TestDocxInsertion(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.docx = str(FIXTURES / "sample_doc.docx")

    def test_insert_at_sentence_end(self):
        doc = Document(self.docx)
        plan = [{"paragraph_id": 1, "sentence_index": 0,
                 "marker": "(Lee et al., 2023)", "claim_id": "C001",
                 "paper_id": "P1", "existing": "keep"}]
        reports = apply_insertion_plan(doc, plan, dry_run=False)
        self.assertEqual(reports[0]["status"], "inserted")
        self.assertIn("(Lee et al., 2023)",
                      doc.paragraphs[1].text)

    def test_no_duplicate_insertion(self):
        doc = Document(self.docx)
        plan = [{"paragraph_id": 4, "sentence_index": 0,
                 "marker": "(Vaswani et al., 2017)", "existing": "keep"}]
        reports = apply_insertion_plan(doc, plan, dry_run=False)
        # paragraph 4 already contains "(Vaswani et al., 2017)"
        self.assertEqual(reports[0]["status"], "skipped_existing")
        text = doc.paragraphs[4].text
        self.assertEqual(text.count("Vaswani"), 1)

    def test_supplement_existing(self):
        doc = Document(self.docx)
        plan = [{"paragraph_id": 4, "sentence_index": 0,
                 "marker": "(Lee et al., 2023)", "existing": "supplement"}]
        reports = apply_insertion_plan(doc, plan, dry_run=False)
        self.assertEqual(reports[0]["status"], "supplemented")

    def test_flag_existing(self):
        doc = Document(self.docx)
        plan = [{"paragraph_id": 4, "sentence_index": 0,
                 "marker": "(X, 2020)", "existing": "flag"}]
        reports = apply_insertion_plan(doc, plan, dry_run=False)
        self.assertEqual(reports[0]["status"], "flag")

    def test_run_level_preserved(self):
        doc = Document(self.docx)
        para = doc.paragraphs[1]
        runs_before = len(para.runs)
        plan = [{"paragraph_id": 1, "sentence_index": 0,
                 "marker": "(Zhao et al., 2022)", "existing": "keep"}]
        apply_insertion_plan(doc, plan, dry_run=False)
        self.assertEqual(len(para.runs), runs_before)  # no new runs created
        self.assertIn("(Zhao et al., 2022)", para.text)


class TestLatex(unittest.TestCase):
    def test_parse(self):
        tex = Path(FIXTURES / "sample.tex").read_text(encoding="utf-8")
        parsed = parse_latex(tex)
        self.assertIn("vaswani2017attention", parsed["existing_keys"])
        self.assertEqual(parsed["bibliography_type"], "bibtex")

    def test_insert_no_duplicate_key(self):
        tex = Path(FIXTURES / "sample.tex").read_text(encoding="utf-8")
        plan = [{"section": "Introduction", "sentence_index": 1,
                 "key": "vaswani2017attention"}]
        new_tex, reports = latex_apply(tex, plan, "cite", dry_run=False)
        self.assertEqual(reports[0]["status"], "skipped_existing")

    def test_insert_new_key(self):
        tex = Path(FIXTURES / "sample.tex").read_text(encoding="utf-8")
        plan = [{"section": "Introduction", "sentence_index": 1,
                 "key": "venkatesh2012utaut2"}]
        new_tex, reports = latex_apply(tex, plan, "parencite", dry_run=False)
        self.assertEqual(reports[0]["status"], "inserted")
        self.assertIn("\\parencite{venkatesh2012utaut2}", new_tex)


class TestBib(unittest.TestCase):
    def test_extract_dois(self):
        text = "doi = {10.1000/abc},\n@article{x, doi = {10.1000/ABC}}"
        dois = extract_dois_from_bib(text)
        self.assertEqual(len(dois), 1)  # case-insensitive dedupe

    def test_apa_order(self):
        papers = [
            {"authors": ["Vaswani, Ashish"], "year": 2017},
            {"authors": ["LeCun, Yann"], "year": 2015},
        ]
        ordered = apa_reference_order(papers)
        self.assertEqual(ordered[0]["authors"][0], "LeCun, Yann")


class TestAudit(unittest.TestCase):
    def test_8B_strips_support_scores(self):
        claims = [{"claim_id": "C1",
                   "original_text": "Generative AI improves learning."}]
        links = [{"claim_id": "C1", "paper_id": "P1",
                  "support_score": 0.95, "support_level": "DIRECT",
                  "evidence_text": ""}]
        papers = [{"paper_id": "P1", "title": "Study", "evidence_text": ""}]
        report = audit_links(claims, links, papers, dry_run_llm=True)
        for link in report["links"]:
            self.assertNotIn("support_score", link)
            self.assertNotIn("support_level", link)

    def test_heuristic_no_evidence_fails(self):
        v = _heuristic_verdict("Some claim", "")
        self.assertEqual(v["verdict"], "FAIL")

    def test_8A_verification_gap_fails(self):
        papers = [{"paper_id": "P1", "title": "Study A", "doi": "10.1/x"}]
        report = audit_document_integrity("", papers, [], {})
        self.assertEqual(report["overall"], "FAIL")
        self.assertIn("verification_log", report["findings"][0]["check"])

    def test_8A_orphan_reference_doi(self):
        papers = [{"paper_id": "P1", "title": "Study A",
                   "doi": "10.1000/real"}]
        vlog = {"P1": {"verdict": "VERIFIED"}}
        report = audit_document_integrity("", papers,
                                          ["https://doi.org/10.9999/orphan"],
                                          vlog)
        self.assertEqual(report["overall"], "FAIL")


def audit_8B(claims, papers, links, **kw):
    return audit_links(claims, links, papers, **kw)


if __name__ == "__main__":
    unittest.main()
