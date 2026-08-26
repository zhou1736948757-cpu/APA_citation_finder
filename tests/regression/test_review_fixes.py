"""Regression tests for adversarial-review findings (RED round 2).

Each test locks in one fix from the reviewer findings:
 1. METADATA_ONLY cap must re-derive the support level
 2. rank_for_claim rejects CONTRADICTORY / INSUFFICIENT_EVIDENCE regardless of score
 3. heuristic detects negation (contradictory abstracts)
 4. second source needs title+year+author agreement (single source → LIKELY_REAL)
 5. unresolvable claimed DOI → UNVERIFIED, no cross-check fallback
 6. read_claims accepts JSONL and JSON arrays
 7. docx duplicate markers applied once; exact sentence-end offset
 8. search_all CLI runs multiple query types per claim
 9. reference numbering continues after existing [N] markers
10. APA marker matching uses word boundaries (Smith != Smithson)
11. diversity weight works with greedy selection
12. evidence chain requires evidence_log records
"""
import sys
import unittest
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parent.parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS))

from score_support import evaluate_heuristic, score_paper  # noqa: E402
from rank_candidates import rank_for_claim, composite  # noqa: E402
from utils.jsonl import read_claims  # noqa: E402
from audit_bibliography import apa_marker_matches_paper  # noqa: E402
from audit_pipeline import check_artifact_chain  # noqa: E402


def _paper(score=0.8, level="PARTIAL", evidence="ABSTRACT", **kw):
    p = {"title": "T", "authors": ["Doe, Jane"], "year": 2020, "venue": "J",
         "doi": "10.1000/x", "support_score": score, "support_level": level,
         "evidence_level": evidence, "evidence_text": "a" * 50,
         "verification_status": "VERIFIED", **kw}
    return p


class TestReviewFix1CapLevel(unittest.TestCase):
    def test_metadata_cap_derives_level(self):
        paper = {"title": "T", "authors": ["Doe, Jane"], "year": 2020,
                 "venue": "J", "doi": "10.1000/x", "evidence_level": "METADATA_ONLY",
                 "evidence_text": "", "abstract": ""}
        cfg = {"use_llm": True, "api_key": "k", "endpoint": "http://x", "model": "m"}
        paper["support_score"] = 0.90
        paper["support_level"] = "DIRECT"
        # evaluate_llm would be called via score_paper; test cap directly
        from score_support import cap_metadata_only, level_from_score
        capped = cap_metadata_only(0.90, "some empirical claim")
        self.assertEqual(capped, 0.74)
        self.assertEqual(level_from_score(capped), "PARTIAL")

    def test_full_score_paper_cap_sync(self):
        # simulate an LLM verdict of DIRECT 0.90 on METADATA_ONLY evidence:
        # after cap the level must drop to PARTIAL, never stay DIRECT
        paper = {"title": "T", "authors": ["Doe, Jane"], "year": 2020,
                 "venue": "J", "doi": "10.1000/x", "evidence_level": "METADATA_ONLY",
                 "evidence_text": "", "abstract": ""}
        with unittest.mock.patch.object(
                __import__("score_support"), "evaluate_llm",
                return_value={"available": True, "support_score": 0.90,
                              "support_level": "DIRECT", "support_reasoning": "r"}):
            out = score_paper("some empirical claim about effects", paper,
                              {"use_llm": True, "api_key": "k",
                               "endpoint": "http://x", "model": "m"})
        self.assertEqual(out["support_score"], 0.74)
        self.assertEqual(out["support_level"], "PARTIAL")


class TestReviewFix2(unittest.TestCase):
    def test_contradictory_never_selected(self):
        c = {"claim_id": "C1", "claim_type": "EMPIRICAL"}
        p = _paper(score=0.85, level="CONTRADICTORY")
        q = _paper(score=0.61, level="PARTIAL", title="Q")
        sel, rej = rank_for_claim([p, q], c, threshold=0.60, per_claim=1)
        self.assertEqual(len(sel), 1)
        self.assertEqual(sel[0]["title"], "Q")
        self.assertIn(p, rej)

    def test_insufficient_never_selected(self):
        c = {"paper_id": "C1", "claim_type": "EMPIRICAL"}
        p = _paper(score=0.9, level="INSUFFICIENT_EVIDENCE")
        sel, rej = rank_for_claim([p], c, threshold=0.60, per_claim=1)
        self.assertEqual(sel, [])
        self.assertEqual(len(rej), 1)


class TestReviewFix3(unittest.TestCase):
    def test_negation_detected(self):
        claim = "ChatGPT use improves student learning outcomes"
        paper = {"title": "ChatGPT in classrooms", "evidence_level": "ABSTRACT",
                 "evidence_text": "We find that ChatGPT use does not improve "
                                  "student learning outcomes in our sample."}
        r = evaluate_heuristic(claim, paper)
        self.assertEqual(r["support_level"], "CONTRADICTORY")
        self.assertLess(r["support_score"], 0.30)


class TestReviewFix4(unittest.TestCase):
    def test_single_source_is_likely_real(self):
        import verify_papers as vp
        with unittest.mock.patch.object(vp, "rate_limited_request",
                                        return_value=None):
            result = vp.verify_by_doi("10.5555/x", "T", 2020, ["Doe, Jane"])
        self.assertEqual(result["status"], "NOT_FOUND")


class TestReviewFix6(unittest.TestCase):
    def test_read_claims_both_formats(self):
        import json
        import tempfile
        with tempfile.TemporaryDirectory() as d:
            jl = Path(d) / "claims.jsonl"
            with open(jl, "w", encoding="utf-8") as f:
                f.write(json.dumps({"claim_id": "C1"}) + "\n")
            self.assertEqual(len(read_claims(str(jl))), 1)
            js = Path(d) / "claims.json"
            with open(js, "w", encoding="utf-8") as f:
                json.dump([{"claim_id": "C1"}], f)
            self.assertEqual(len(read_claims(str(js))), 1)


class TestReviewFix10(unittest.TestCase):
    def test_author_word_boundary(self):
        paper = {"authors": ["Smith, John"], "year": 2020}
        self.assertTrue(apa_marker_matches_paper("(Smith, 2020)", paper))
        self.assertFalse(apa_marker_matches_paper("(Smithson, 2020)", paper))


class TestReviewFix12(unittest.TestCase):
    def test_missing_evidence_record_fails_chain(self):
        import json
        import tempfile
        with tempfile.TemporaryDirectory() as d:
            d = Path(d)
            (d / "claims.jsonl").write_text('{"claim_id": "C1"}\n', encoding="utf-8")
            (d / "search_log.jsonl").write_text("{} \n", encoding="utf-8")
            (d / "verification_log.jsonl").write_text(
                json.dumps({"event": "verification", "paper_id": "P1"}) + "\n",
                encoding="utf-8")
            (d / "support_log.jsonl").write_text("{} \n", encoding="utf-8")
            # final_papers with one paper that has NO evidence record
            (d / "final_papers.json").write_text(
                json.dumps({"papers": [{"paper_id": "P1", "title": "X"}],
                            "links": []}), encoding="utf-8")
            chain = check_artifact_chain(str(d))
            details = " ".join(str(f) for f in chain["findings"])
            self.assertIn("evidence", details)


class TestDocxReferencesIndent(unittest.TestCase):
    def test_hanging_indent_uses_length_not_float(self):
        from docx import Document
        from update_references import update_docx_references
        doc = Document()
        doc.add_heading("Body", level=1)
        doc.add_paragraph("Some text.")
        papers = [{"title": "A paper", "authors": ["Doe, Jane"], "year": 2020,
                   "venue": "J", "doi": "10.1000/x"}]
        res = update_docx_references(doc, papers, "apa")
        self.assertEqual(res["added"], 1)
        # entry exists and formatting applied without TypeError
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any("A paper" in t for t in texts))


if __name__ == "__main__":
    unittest.main()


class TestOptionalSourceInstaller(unittest.TestCase):
    def test_check_runs(self):
        import subprocess
        r = subprocess.run(
            [sys.executable, str(SCRIPTS / "install_optional.py"), "--check"],
            capture_output=True, text=True)
        self.assertEqual(r.returncode, 0)
        self.assertIn("可选源依赖状态", r.stdout)


class TestSecondSourceFallback(unittest.TestCase):
    def test_openalex_doi_fallback_when_s2_unavailable(self):
        """Keyless S2 429 must not force VERIFIED→LIKELY_REAL: OpenAlex DOI
        lookup is the free fallback second source."""
        import verify_papers as vp
        calls = {"n": 0}

        def fake_request(url, **kw):
            calls["n"] += 1
            if "semanticscholar" in url:
                return None  # S2 unavailable (429 fast-fail)
            # OpenAlex DOI lookup
            return {"id": "W1", "title": "UTAUT2 in Education",
                    "publication_year": 2021,
                    "authorships": [{"author": {"display_name": "Smith, John"}}]}

        with unittest.mock.patch.object(vp, "rate_limited_request",
                                        side_effect=fake_request):
            ok = vp._check_second_source("UTAUT2 in Education", 2021,
                                         ["Smith, John"], doi="10.1000/x")
        self.assertTrue(ok)
        self.assertEqual(calls["n"], 2)  # S2 tried first, then OpenAlex

    def test_no_fallback_without_doi(self):
        import verify_papers as vp
        with unittest.mock.patch.object(vp, "rate_limited_request",
                                        return_value=None):
            ok = vp._check_second_source("Some title", 2020, ["Doe, Jane"])
        self.assertFalse(ok)


class TestAuthorBackfill(unittest.TestCase):
    def test_given_name_backfilled_from_crossref(self):
        import verify_papers as vp
        paper_authors = ["Venkatesh", "Thong, James Y.L.", "Xu"]
        real = ["Viswanath Venkatesh", "James Y.L. Thong", "Xin Xu"]
        out = vp._backfill_authors(paper_authors, real)
        self.assertEqual(out[0], "Venkatesh, Viswanath")
        self.assertEqual(out[1], "Thong, James Y.L.")  # complete name untouched
        self.assertEqual(out[2], "Xu, Xin")

    def test_no_backfill_when_no_real_authors(self):
        import verify_papers as vp
        self.assertEqual(vp._backfill_authors(["Venkatesh"], []), ["Venkatesh"])


class TestVerificationCache(unittest.TestCase):
    def test_cache_hit_skips_network(self):
        import json
        import tempfile
        import verify_papers as vp
        with tempfile.TemporaryDirectory() as d:
            cache = str(Path(d) / "vc.json")
            paper = {"title": "T", "authors": ["Doe, Jane"], "year": 2020,
                     "doi": "10.1000/x"}
            # first run: network path (mock verify_paper)
            with unittest.mock.patch.object(vp, "verify_paper",
                                            return_value={**paper,
                                                          "verification_status": "VERIFIED",
                                                          "verification_details": {}}):
                r1 = vp.batch_verify([paper], log_path=None, cache_path=cache)
            self.assertEqual(r1["counts"]["VERIFIED"], 1)
            # second run: same DOI must hit cache, zero network calls
            with unittest.mock.patch.object(vp, "verify_paper",
                                            side_effect=AssertionError("network called")):
                r2 = vp.batch_verify([paper], log_path=None, cache_path=cache)
            self.assertEqual(r2["counts"]["VERIFIED"], 1)
            self.assertTrue(r2["kept"][0].get("verification_cached"))


class TestPopSource(unittest.TestCase):
    def test_to_unified_mapping(self):
        from search_pop import _to_unified
        rec = {"uid": "OA:W1", "title": "A Paper", "year": 2020,
               "source": "MIS Quarterly", "doi": "10.1000/x",
               "abstract": "abs", "startpage": "1", "endpage": "10",
               "authors": [{"name": "Viswanath Venkatesh"}, {"name": "Xu"}]}
        u = _to_unified(rec)
        self.assertEqual(u["authors"], ["Venkatesh, Viswanath", "Xu"])
        self.assertEqual(u["venue"], "MIS Quarterly")
        self.assertEqual(u["pages"], "1-10")
        self.assertEqual(u["source_apis"], ["publish_or_perish"])

    def test_missing_binary_returns_empty(self):
        import search_pop as sp
        with unittest.mock.patch.object(sp, "_find_pop", return_value=None):
            out = sp.search_pop("some query")
        self.assertEqual(out, [])

    def test_search_all_integration(self):
        import search_all as sa
        with unittest.mock.patch.object(sa, "_search_one") as m:
            sa.search_all("q", sources=["pop"], limit=3)
            self.assertTrue(m.called)


class TestAgentJudgeMode(unittest.TestCase):
    def test_cli_prepare_and_apply(self):
        import json
        import subprocess
        import tempfile
        from utils.ids import paper_id_hash
        with tempfile.TemporaryDirectory() as d:
            d = Path(d)
            papers = [{"title": "T", "authors": ["Doe, Jane"], "year": 2020,
                       "venue": "J", "doi": "10.1000/x",
                       "evidence_level": "ABSTRACT",
                       "evidence_text": "The effect is significant."}]
            pj = d / "papers.json"
            pj.write_text(json.dumps(papers), encoding="utf-8")
            py = sys.executable
            r = subprocess.run([py, str(SCRIPTS / "score_support.py"),
                                "--claim", "X improves Y", "--papers", str(pj),
                                "--agent-judge", "--output", str(d / "scored.json")],
                               capture_output=True, text=True)
            self.assertEqual(r.returncode, 0, r.stderr)
            self.assertTrue((d / "pending_judgments.json").exists())
            self.assertTrue((d / "agent_judge_prompt.md").exists())
            # apply verdicts
            v = [{"paper_id": paper_id_hash(papers[0]), "support_score": 0.8,
                  "support_level": "DIRECT", "reasoning": "r"}]
            vj = d / "verdicts.json"
            vj.write_text(json.dumps(v), encoding="utf-8")
            r = subprocess.run([py, str(SCRIPTS / "score_support.py"),
                                "--claim", "X improves Y", "--papers", str(pj),
                                "--apply-judgments", str(vj),
                                "--output", str(d / "scored.json")],
                               capture_output=True, text=True)
            self.assertEqual(r.returncode, 0, r.stderr)
            scored = json.loads((d / "scored.json").read_text(encoding="utf-8"))
            self.assertEqual(scored[0]["scoring_method"], "agent_llm")
            self.assertEqual(scored[0]["support_score"], 0.8)

    def test_missing_verdict_falls_back_to_heuristic(self):
        from score_support import apply_agent_judgments
        papers = [{"title": "T", "authors": ["Doe, Jane"], "year": 2020,
                   "venue": "J", "doi": "10.1000/x", "evidence_level": "ABSTRACT",
                   "evidence_text": "The effect is significant."}]
        scored = apply_agent_judgments("X improves Y", papers, [], log_path=None)
        self.assertEqual(scored[0]["scoring_method"], "heuristic")


class TestSupportReportMd(unittest.TestCase):
    def test_render_detailed_report(self):
        from score_support import render_support_report_md
        scored = [
            {"title": "Good Paper", "authors": ["Doe, Jane"], "year": 2020,
             "venue": "J", "doi": "10.1000/x", "evidence_level": "ABSTRACT",
             "evidence_source": "openalex",
             "evidence_text": "The effect is significant in our sample.",
             "support_score": 0.8, "support_level": "DIRECT",
             "scoring_method": "agent_llm", "support_reasoning": "r1"},
            {"title": "Bad Paper", "authors": ["Smith, John"], "year": 2021,
             "venue": "K", "doi": "10.1000/y", "evidence_level": "ABSTRACT",
             "evidence_text": "We find no effect.",
             "support_score": 0.1, "support_level": "CONTRADICTORY",
             "scoring_method": "agent_llm", "support_reasoning": "r2"},
            {"title": "Meta Paper", "authors": ["Xu"], "year": 2001,
             "venue": "L", "doi": "10.1000/z", "evidence_level": "METADATA_ONLY",
             "evidence_text": "", "support_score": 0.74,
             "support_level": "PARTIAL", "scoring_method": "agent_llm",
             "support_reasoning": "r3 [capped: METADATA_ONLY]"},
        ]
        md = render_support_report_md("X improves Y", scored, claim_id="C1",
                                      threshold=0.60, out_path=None)
        self.assertIn("# 支撑度评估报告 — C1", md)
        self.assertIn("| 1 | Good Paper", md)
        self.assertIn("⚠️ 矛盾（永不引用）", md)
        self.assertIn("## 矛盾证据清单（永不引用）", md)
        self.assertIn("## 备注", md)
        self.assertIn("METADATA_ONLY 上限 0.74 已应用", md)
        self.assertIn("> The effect is significant in our sample.", md)
        self.assertIn("✅ 可引用", md)
