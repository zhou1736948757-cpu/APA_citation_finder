"""APA_citation_finder :: insert_docx.py
Stage 8c — Word (.docx) citation insertion.

Design rules (spec §37/§38):
  * NO global string replace. Insertion happens at python-docx run level so
    fonts/formatting are preserved; only the run's text changes.
  * Existing citation markers are parsed first:
        (Author, 2020) | (A & B, 2021) | (X et al., 2022) | [1], [2-4]
    → decision per target sentence: keep / supplement / replace / flag
  * Never insert the same citation twice; never renumber existing [N] refs.
  * Tables/footnotes/headings are untouched.

Input: insertion plan JSON (agent-generated from final_papers links):
  {
    "style": "apa|ieee|nature|vancouver|gb-t-7714",
    "insertions": [
      {
        "paragraph_id": 6,
        "sentence_index": 2,          // 0-based sentence within paragraph
        "marker": "(Vaswani et al., 2017)",
        "claim_id": "C001",
        "paper_id": "P001",
        "existing": "keep|supplement|flag"     // optional decision
      }
    ]
  }

Usage:
  python insert_docx.py in.docx plan.json --output out.docx
  python insert_docx.py in.docx plan.json --dry-run
"""
from __future__ import annotations

import argparse
import json
import sys
from typing import Any

from utils.jsonl import load_json
from utils.text import split_sentences_with_offsets
from utils.docx_helpers import (
    find_existing_citation_spans,
    insert_marker_at_offset,
    insert_marker_before_punctuation,
)

try:
    from docx import Document
except ImportError as e:
    print("[APA_citation_finder] python-docx required: pip install python-docx", file=sys.stderr)
    raise


def locate_sentence_run(para, sentence_index: int) -> dict:
    """Return {run, offset_in_run, sentence_text} — the exact char offset
    inside the run where the sentence's final punctuation starts.

    Using the offset (not just the run) matters when one run spans several
    sentences: the marker must land after THIS sentence, not at the run end.
    """
    text = para.text or ""
    sents = split_sentences_with_offsets(text)
    if sentence_index >= len(sents):
        raise LookupError(f"sentence {sentence_index} not found in paragraph")
    sent = sents[sentence_index]
    end_char = sent["end"]

    # end_char points just past the sentence; walk back over trailing
    # punctuation so the marker is inserted BEFORE the final period
    sent_text = sent["text"]
    punct_len = 0
    while punct_len < len(sent_text) and sent_text[-1 - punct_len] in ".!?。！？":
        punct_len += 1
    insert_char = end_char - punct_len

    acc = 0
    for run in para.runs:
        run_len = len(run.text or "")
        if insert_char <= acc + run_len:
            return {"run": run, "offset_in_run": insert_char - acc,
                    "sentence_text": sent["text"]}
        acc += run_len
    if para.runs:
        return {"run": para.runs[-1], "offset_in_run": len(para.runs[-1].text or ""),
                "sentence_text": sent["text"]}
    raise LookupError("paragraph has no runs")


def apply_insertion_plan(doc, plan: list[dict], dry_run: bool = False) -> list[dict]:
    """Apply markers at run level with exact sentence-end offsets.

    Returns per-insertion reports. Duplicate markers for the same
    (paragraph, sentence, marker) are applied at most once.
    """
    reports: list[dict] = []
    applied: set[tuple] = set()
    for ins in plan:
        para_idx = ins.get("paragraph_id")
        if para_idx is None or para_idx >= len(doc.paragraphs):
            reports.append({**ins, "status": "skip", "reason": "paragraph out of range"})
            continue
        para = doc.paragraphs[para_idx]
        sent_idx = ins.get("sentence_index", 0)
        marker = ins.get("marker", "")

        # cross-insertion dedup: identical marker at same spot inserted once
        dedup_key = (para_idx, sent_idx, marker)
        if dedup_key in applied:
            reports.append({**ins, "status": "skipped_existing",
                            "reason": "duplicate marker already applied"})
            continue

        para_text = para.text or ""
        spans = find_existing_citation_spans(para_text)
        if spans:
            decision = ins.get("existing", "keep")
            if decision == "flag":
                reports.append({**ins, "status": "flag",
                                "reason": "existing citation present; human review"})
                continue
            if decision == "supplement":
                try:
                    loc = locate_sentence_run(para, sent_idx)
                except LookupError as e:
                    reports.append({**ins, "status": "skip", "reason": str(e)})
                    continue
                if not dry_run:
                    insert_marker_at_offset(loc["run"], loc["offset_in_run"], marker)
                    applied.add(dedup_key)
                reports.append({**ins, "status": "supplemented"})
                continue
            # keep → do not duplicate
            reports.append({**ins, "status": "skipped_existing",
                            "reason": "target sentence already has a citation"})
            continue

        try:
            loc = locate_sentence_run(para, sent_idx)
        except LookupError as e:
            reports.append({**ins, "status": "skip", "reason": str(e)})
            continue
        if not dry_run:
            insert_marker_at_offset(loc["run"], loc["offset_in_run"], marker)
            applied.add(dedup_key)
        reports.append({**ins, "status": "inserted"})
    return reports


def _cli() -> int:
    p = argparse.ArgumentParser(description="APA_citation_finder Word citation inserter")
    p.add_argument("docx_file")
    p.add_argument("plan_json", help="insertion plan JSON (list or {insertions:[...]})")
    p.add_argument("--output")
    p.add_argument("--dry-run", action="store_true")
    args = p.parse_args()

    plan_data = load_json(args.plan_json)
    plan = plan_data.get("insertions") if isinstance(plan_data, dict) else plan_data

    doc = Document(args.docx_file)
    reports = apply_insertion_plan(doc, plan, dry_run=args.dry_run)

    if args.dry_run:
        print(json.dumps(reports, ensure_ascii=False, indent=2))
        return 0

    out_path = args.output or args.docx_file.replace(".docx", "_cited.docx")
    doc.save(out_path)
    counts: dict[str, int] = {}
    for r in reports:
        counts[r.get("status", "?")] = counts.get(r.get("status", "?"), 0) + 1
    print(json.dumps({"output": out_path, "reports": reports, "counts": counts},
                     ensure_ascii=False, indent=2))
    return 0 if reports else 1


if __name__ == "__main__":
    sys.exit(_cli())
