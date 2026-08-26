"""APA_citation_finder :: update_references.py
Stage 8e — Bibliography update.

Modes:
  * docx            — append formatted entries to an existing References /
                      Bibliography section, or create one at document end.
                      APA: alphabetical; numeric styles: first-appearance order.
  * latex-bibtex    — append missing entries to a .bib file (DOI-deduped).
  * latex-thebib    — insert \\bibitem entries into \\begin{thebibliography}.

Never duplicates: same DOI / same bibtex key / same title already present.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from format_citation import (
    format_citation,
    format_bibtex_entry,
    apa_reference_list,
)
from utils.ids import normalize_doi
from utils.jsonl import load_json, write_json

BIB_ITEM_RE = re.compile(r"\\bibitem[^}]*\}", re.S)


def apa_reference_order(papers: list[dict]) -> list[dict]:
    """APA: sort by first-author last name, then year."""
    def key(p: dict):
        authors = p.get("authors") or []
        first = authors[0] if authors else ""
        last = first.split(",")[0].strip().lower() if "," in first else first.lower()
        return (last, p.get("year") or 9999)
    return sorted(papers, key=key)


def _sorted_for_style(papers: list[dict], style: str) -> list[dict]:
    if style in ("apa", "apa7"):
        return apa_reference_order(papers)
    return papers  # numeric styles: caller provides first-appearance order


def _sorted_style_entries(papers: list[dict], style: str,
                          start_number: int = 1) -> list[str]:
    """Format entries; numeric styles number from start_number (so appends
    continue after existing [N] markers instead of colliding with them)."""
    ordered = _sorted_for_style(papers, style)
    numbered = style in ("ieee", "nature", "vancouver", "gb-t-7714", "gbt7714", "gb-t")
    if not numbered:
        return [format_citation(p, style, None) for p in ordered]
    return [format_citation(p, style, start_number + i)
            for i, p in enumerate(ordered)]


def update_docx_references(doc, papers: list[dict], style: str,
                           section_name: str = "References") -> dict:
    """Update or create a References section in an opened docx Document."""
    from utils.docx_helpers import APA_KEYWORDS

    # locate existing references section
    ref_start = None
    ref_end = None
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if text and any(k in text.lower() for k in ("reference", "bibliography", "参考文献")):
            if ref_start is None:
                ref_start = i
    if ref_start is None:
        # create new section at the end
        doc.add_heading(section_name, level=1)
        ref_start = len(doc.paragraphs) - 1
    # find existing entries (paragraphs after the heading)
    existing_texts = [p.text.strip() for p in doc.paragraphs[ref_start + 1:]
                      if (p.text or "").strip()]
    existing_dois = set()
    for t in existing_texts:
        m = re.search(r"doi\.org/([^\s\)\]]+)", t)
        if m:
            existing_dois.add(m.group(1).lower())

    start_num = 1
    if style in ("ieee", "nature", "vancouver", "gb-t-7714", "gbt7714", "gb-t"):
        nums = [int(n) for n in NUMERIC_DOC_RE.findall(
            "\n".join((p.text or "") for p in doc.paragraphs))]
        if nums:
            start_num = max(nums) + 1
    entries = _sorted_style_entries(papers, style, start_number=start_num)
    added = 0
    skipped = 0
    for entry in entries:
        m = re.search(r"doi\.org/([^\s\)\]]+)", entry)
        doi = m.group(1).lower() if m else None
        if doi and doi in existing_dois:
            skipped += 1
            continue
        if any(entry.split("(")[0].strip()[:60] in t for t in existing_texts):
            skipped += 1
            continue
        para = doc.add_paragraph(entry)
        # hanging indent (APA); python-docx needs Length objects, not floats
        from docx.shared import Inches
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        added += 1
    return {"added": added, "skipped": skipped, "section": section_name}


def update_latex_thebibliography(tex: str, papers: list[dict], style: str) -> str:
    """Insert \\bibitem entries before \\end{thebibliography}."""
    m = re.search(r"\\begin\{thebibliography\}(\{.*?\})?(.*?)\\end\{thebibliography\}",
                  tex, re.S)
    if not m:
        return tex
    existing = set(BIB_ITEM_RE.findall(m.group(2)))
    # continue numbering after the highest existing \bibitem index
    idx = [int(k) for k in existing if k.isdigit()]
    start_number = (max(idx) + 1) if idx else 1
    new_items = []
    for n, p in enumerate(papers):
        entry = format_citation(p, style if style != "apa7" else "apa7",
                                start_number + n)
        key = p.get("bibtex_key") or p.get("paper_id") or "paper"
        if key in existing:
            continue
        item = f"\\bibitem{{{key}}} {entry}"
        new_items.append(item)
        existing.add(key)
    if not new_items:
        return tex
    insertion = "\n" + "\n".join(new_items) + "\n"
    return tex[: m.end(2)] + insertion + tex[m.end(2):]


def update_bibtex_file(bib_path: str, papers: list[dict]) -> dict:
    """Append missing entries to an existing .bib (DOI/key dedup)."""
    from format_bibtex import extract_dois_from_bib, batch_to_bibtex
    bib = Path(bib_path)
    existing_text = bib.read_text(encoding="utf-8") if bib.exists() else ""
    existing_dois = extract_dois_from_bib(existing_text)
    new = batch_to_bibtex(papers, existing_dois)
    if not new:
        return {"added": 0, "skipped": len(papers), "path": str(bib)}
    with open(bib_path, "a", encoding="utf-8") as f:
        if existing_text and not existing_text.endswith("\n"):
            f.write("\n")
        f.write(new)
    added = new.count("@")
    return {"added": added, "skipped": len(papers) - added, "path": str(bib)}


def _cli() -> int:
    p = argparse.ArgumentParser(description="APA_citation_finder bibliography updater")
    p.add_argument("--mode", required=True, choices=["docx", "thebibliography", "bibtex"])
    p.add_argument("--papers", required=True, help="final_papers.json (list or envelope)")
    p.add_argument("--style", default="apa")
    p.add_argument("--docx", help="target .docx (mode=docx)")
    p.add_argument("--tex", help="target .tex (mode=thebibliography)")
    p.add_argument("--bib", help="target .bib (mode=bibtex)")
    p.add_argument("--output")
    args = p.parse_args()

    data = load_json(args.papers)
    papers = data.get("papers") if isinstance(data, dict) and "papers" in data else (
        data if isinstance(data, list) else [])

    if args.mode == "docx":
        from docx import Document
        if not args.docx:
            p.error("--docx required")
        doc = Document(args.docx)
        result = update_docx_references(doc, papers, args.style)
        out = args.output or args.docx
        doc.save(out)
        result["output"] = out
    elif args.mode == "thebibliography":
        if not args.tex:
            p.error("--tex required")
        tex = Path(args.tex).read_text(encoding="utf-8")
        new_tex = update_latex_thebibliography(tex, papers, args.style)
        out = args.output or args.tex
        Path(out).write_text(new_tex, encoding="utf-8")
        result = {"output": out, "entries": len(papers)}
    else:
        if not args.bib:
            p.error("--bib required")
        result = update_bibtex_file(args.bib, papers)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
